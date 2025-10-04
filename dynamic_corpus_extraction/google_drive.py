"""Google Drive PDF integration for the dynamic corpus extraction engine."""

from __future__ import annotations

import importlib
import importlib.util
import io
import json
import re
from typing import Callable, Iterable, Iterator, Mapping, MutableMapping, Sequence
from urllib.error import HTTPError, URLError
from urllib.parse import parse_qsl, urlencode, urlparse, urlunparse
from urllib.request import Request, build_opener

from .engine import CorpusExtractionContext, ExtractionLoader

__all__ = [
    "GoogleDriveClient",
    "build_google_drive_pdf_loader",
    "parse_drive_share_link",
]

_PDF_MIME_TYPE = "application/pdf"
_FOLDER_MIME_TYPE = "application/vnd.google-apps.folder"
_DOCX_MIME_TYPES = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
)
_DEFAULT_FIELDS = "nextPageToken, files(id, name, mimeType, modifiedTime, size, md5Checksum, webViewLink, parents)"


def _append_query(url: str, params: Mapping[str, str] | None) -> str:
    if not params:
        return url
    parsed = urlparse(url)
    current = dict(parse_qsl(parsed.query, keep_blank_values=True))
    for key, value in params.items():
        if value is not None:
            current[key] = value
    encoded = urlencode(current, doseq=True)
    return urlunparse(parsed._replace(query=encoded))


def _load_pypdf() -> "PyPDF2":  # type: ignore[name-defined]
    try:
        import PyPDF2  # type: ignore[import]
    except ModuleNotFoundError as exc:  # pragma: no cover - import guard
        raise RuntimeError(
            "PyPDF2 is required for PDF text extraction. Install it with `pip install PyPDF2`."
        ) from exc
    return PyPDF2


def _default_pdf_text_extractor(payload: bytes, *, file_name: str) -> str:
    PyPDF2 = _load_pypdf()
    reader = PyPDF2.PdfReader(io.BytesIO(payload))
    text_parts: list[str] = []
    for page in reader.pages:
        extracted = page.extract_text() or ""
        if extracted:
            text_parts.append(extracted)
    text = "\n".join(part.strip() for part in text_parts if part.strip())
    if not text:
        raise RuntimeError(f"No extractable text found in PDF '{file_name}'")
    return text


def _load_python_docx() -> "docx":  # type: ignore[name-defined]
    try:
        import docx  # type: ignore[import]
    except ModuleNotFoundError as exc:  # pragma: no cover - import guard
        raise RuntimeError(
            "python-docx is required for DOCX text extraction. Install it with `pip install python-docx`."
        ) from exc
    return docx


def _default_docx_text_extractor(payload: bytes, *, file_name: str) -> str:
    package = _load_python_docx()
    try:
        document = package.Document(io.BytesIO(payload))
    except Exception as error:  # pragma: no cover - defensive guard
        raise RuntimeError(f"Failed to read DOCX '{file_name}'") from error

    text_parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text_parts.append(cell_text)

    text = "\n".join(text_parts).strip()
    if not text:
        raise RuntimeError(f"No extractable text found in DOCX '{file_name}'")
    return text


def _extract_pdf_page_texts(
    payload: bytes,
    *,
    file_name: str,
    enable_ocr: bool,
    languages: str | None,
    dpi: int,
) -> list[str]:
    """Return per-page text for a PDF, optionally filling gaps with OCR."""

    PyPDF2 = _load_pypdf()
    reader = PyPDF2.PdfReader(io.BytesIO(payload))
    page_texts: list[str] = []
    for page_index, page in enumerate(reader.pages, start=1):
        try:
            extracted = page.extract_text() or ""
        except Exception as error:  # pragma: no cover - defensive guard
            raise RuntimeError(
                f"Failed to extract text from page {page_index} of PDF '{file_name}'"
            ) from error
        page_texts.append(extracted.strip())

    if not enable_ocr:
        return page_texts

    if not any(not text for text in page_texts):
        return page_texts

    pdf2image, pytesseract = _load_pdf_ocr_toolchain()
    try:
        images = pdf2image.convert_from_bytes(payload, dpi=dpi)
    except Exception as error:  # pragma: no cover - conversion failure
        raise RuntimeError(f"Failed to rasterise PDF '{file_name}' for OCR") from error

    try:
        for index, image in enumerate(images):
            try:
                if index >= len(page_texts) or page_texts[index]:
                    continue
                if languages:
                    extracted = pytesseract.image_to_string(image, lang=languages)
                else:
                    extracted = pytesseract.image_to_string(image)
                cleaned = extracted.strip()
                if cleaned:
                    page_texts[index] = cleaned
            finally:
                try:
                    image.close()
                except Exception:  # pragma: no cover - defensive cleanup
                    pass
    finally:
        for image in images:
            try:
                image.close()
            except Exception:  # pragma: no cover - defensive cleanup
                pass

    return page_texts


def _normalise_ocr_languages(languages: Sequence[str] | str | None) -> str | None:
    if languages is None:
        return None
    if isinstance(languages, str):
        cleaned = languages.strip()
        return cleaned or None
    cleaned_sequence: list[str] = []
    for candidate in languages:
        text = str(candidate).strip()
        if text:
            cleaned_sequence.append(text)
    unique = []
    seen: set[str] = set()
    for entry in cleaned_sequence:
        if entry and entry not in seen:
            seen.add(entry)
            unique.append(entry)
    if not unique:
        return None
    return "+".join(unique)


def _load_pdf_ocr_toolchain() -> tuple[object, object]:
    dependency_map = (
        ("pdf2image", "pdf2image"),
        ("pytesseract", "pytesseract"),
        ("PIL.Image", "Pillow"),
    )
    for module_name, package_name in dependency_map:
        if importlib.util.find_spec(module_name) is None:
            raise RuntimeError(
                "PDF OCR support requires the "
                f"'{package_name}' package. Install it with `pip install {package_name}`."
            )

    pdf2image = importlib.import_module("pdf2image")
    pytesseract = importlib.import_module("pytesseract")
    importlib.import_module("PIL.Image")
    return pdf2image, pytesseract


def _ocr_pdf_text_extractor(
    payload: bytes,
    *,
    file_name: str,
    languages: str | None,
    dpi: int,
) -> str:
    pdf2image, pytesseract = _load_pdf_ocr_toolchain()
    try:
        images = pdf2image.convert_from_bytes(payload, dpi=dpi)
    except Exception as error:  # pragma: no cover - conversion failure
        raise RuntimeError(f"Failed to rasterise PDF '{file_name}' for OCR") from error

    text_parts: list[str] = []
    try:
        for image in images:
            try:
                if languages:
                    extracted = pytesseract.image_to_string(image, lang=languages)
                else:
                    extracted = pytesseract.image_to_string(image)
            finally:
                try:
                    image.close()
                except Exception:  # pragma: no cover - defensive cleanup
                    pass
            cleaned = extracted.strip()
            if cleaned:
                text_parts.append(cleaned)
    finally:
        for image in images:
            try:
                image.close()
            except Exception:  # pragma: no cover - defensive cleanup
                pass

    text = "\n".join(text_parts)
    if not text.strip():
        raise RuntimeError(f"OCR extraction produced no text for PDF '{file_name}'")
    return text


class GoogleDriveClient:
    """Minimal Google Drive client focused on folder traversal and downloads."""

    def __init__(
        self,
        *,
        api_key: str | None = None,
        access_token: str | None = None,
        base_url: str = "https://www.googleapis.com/drive/v3",
        opener_factory: Callable[[], object] | None = None,
    ) -> None:
        api_key = api_key.strip() if api_key else None
        access_token = access_token.strip() if access_token else None
        if not api_key and not access_token:
            raise ValueError("Either api_key or access_token must be provided")
        self._api_key = api_key
        self._access_token = access_token
        self._base_url = base_url.rstrip("/")
        opener = opener_factory() if opener_factory is not None else build_opener()
        self._opener = opener

    # ----------------------------------------------------------------- utilities
    def _build_headers(self, accept: str) -> Mapping[str, str]:
        headers = {"Accept": accept}
        if self._access_token:
            headers["Authorization"] = f"Bearer {self._access_token}"
        return headers

    def _prepare_params(self, params: Mapping[str, str] | None) -> Mapping[str, str]:
        combined = dict(params or {})
        if self._api_key:
            combined.setdefault("key", self._api_key)
        return combined

    def _request_bytes(
        self,
        path: str,
        *,
        params: Mapping[str, str] | None = None,
        accept: str = "application/json",
    ) -> bytes:
        url = f"{self._base_url}/{path.lstrip('/') }"
        request = Request(_append_query(url, self._prepare_params(params)), headers=self._build_headers(accept))
        try:
            with self._opener.open(request) as response:
                return response.read()
        except HTTPError as error:  # pragma: no cover - network failure surface
            raise RuntimeError(
                f"Google Drive request failed with {error.code} {error.reason}: {error.read().decode('utf-8', errors='ignore')}"
            ) from error
        except URLError as error:  # pragma: no cover - network failure surface
            raise RuntimeError(f"Google Drive request to {url} failed: {error.reason}") from error

    def _request_json(
        self,
        path: str,
        *,
        params: Mapping[str, str] | None = None,
    ) -> MutableMapping[str, object]:
        payload = self._request_bytes(path, params=params, accept="application/json")
        try:
            return json.loads(payload.decode("utf-8"))
        except json.JSONDecodeError as error:
            raise RuntimeError(f"Failed to decode Google Drive JSON response from {path}") from error

    # ----------------------------------------------------------------- operations
    def iter_files(
        self,
        *,
        folder_id: str,
        mime_types: Sequence[str] | None = None,
        page_size: int = 200,
        fields: str = _DEFAULT_FIELDS,
    ) -> Iterator[MutableMapping[str, object]]:
        query_parts = [f"'{folder_id}' in parents", "trashed = false"]
        if mime_types:
            mime_checks = [f"mimeType = '{mime}'" for mime in mime_types]
            query_parts.append("(" + " or ".join(mime_checks) + ")")
        params: dict[str, str] = {
            "q": " and ".join(query_parts),
            "pageSize": str(page_size),
            "supportsAllDrives": "true",
            "includeItemsFromAllDrives": "true",
            "fields": fields,
        }
        next_token: str | None = None
        while True:
            effective_params = dict(params)
            if next_token:
                effective_params["pageToken"] = next_token
            response = self._request_json("files", params=effective_params)
            for entry in response.get("files", []):
                if isinstance(entry, MutableMapping):
                    yield entry
            next_token = response.get("nextPageToken")
            if not next_token:
                break

    def get_file_metadata(self, file_id: str, *, fields: str | None = None) -> MutableMapping[str, object]:
        params = {"supportsAllDrives": "true"}
        if fields:
            params["fields"] = fields
        return self._request_json(f"files/{file_id}", params=params)

    def download_file(self, file_id: str) -> bytes:
        return self._request_bytes(
            f"files/{file_id}",
            params={"alt": "media", "supportsAllDrives": "true"},
            accept="application/pdf, */*;q=0.8",
        )


_FOLDER_PATTERN = re.compile(r"/folders/([A-Za-z0-9_-]+)")
_FILE_PATTERN = re.compile(r"/file/d/([A-Za-z0-9_-]+)")


def parse_drive_share_link(link: str) -> tuple[str, str]:
    """Return (target_type, identifier) from a Google Drive share link."""

    candidate = (link or "").strip()
    if not candidate:
        raise ValueError("share link must not be empty")

    parsed = urlparse(candidate)
    if not parsed.scheme or not parsed.netloc:
        raise ValueError("share link must include scheme and host")

    folder_match = _FOLDER_PATTERN.search(parsed.path)
    if folder_match:
        return "folder", folder_match.group(1)

    file_match = _FILE_PATTERN.search(parsed.path)
    if file_match:
        return "file", file_match.group(1)

    if parsed.query:
        params = dict(parse_qsl(parsed.query, keep_blank_values=True))
        identifier = params.get("id")
        if identifier:
            return "file", identifier

    raise ValueError("Unable to extract Google Drive identifier from share link")


def _coerce_size(value: object) -> int | None:
    if value is None:
        return None
    try:
        return int(value)
    except (TypeError, ValueError):  # pragma: no cover - defensive path
        return None


def _coerce_positive_int(value: object) -> int | None:
    candidate = _coerce_size(value)
    if candidate is None or candidate <= 0:
        return None
    return candidate


def _normalise_identifier_set(values: Sequence[object] | object | None) -> set[str]:
    """Return a set of cleaned identifiers from ``values``."""

    if values is None:
        return set()
    if isinstance(values, str):
        cleaned = values.strip()
        return {cleaned} if cleaned else set()
    try:
        iterator = iter(values)  # type: ignore[arg-type]
    except TypeError:
        single = str(values).strip()
        return {single} if single else set()
    cleaned: set[str] = set()
    for entry in iterator:
        text = str(entry).strip()
        if text:
            cleaned.add(text)
    return cleaned


def build_google_drive_pdf_loader(
    *,
    share_link: str | None = None,
    folder_id: str | None = None,
    file_ids: Sequence[str] | None = None,
    api_key: str | None = None,
    access_token: str | None = None,
    tags: Sequence[str] | None = ("google_drive", "pdf"),
    client_factory: Callable[[], GoogleDriveClient] | None = None,
    pdf_text_extractor: Callable[[bytes, Mapping[str, object]], str] | None = None,
    max_file_size: int | None = 50_000_000,
    batch_size: int | None = None,
    page_batch_size: int | None = None,
    enable_ocr: bool = False,
    ocr_languages: Sequence[str] | str | None = ("eng",),
    ocr_dpi: int = 300,
    include_docx: bool = False,
    docx_text_extractor: Callable[[bytes, Mapping[str, object]], str] | None = None,
    skip_file_ids: Sequence[str] | None = None,
) -> ExtractionLoader:
    """Create a loader that streams Google Drive PDFs as corpus documents.

    When ``enable_ocr`` is ``True`` the loader falls back to rasterising PDF
    pages and running them through Tesseract OCR whenever text extraction
    yields empty results.  This allows the pipeline to capture scanned
    documents while still preferring the faster embedded text path when
    available. Set ``skip_file_ids`` to ignore previously processed Drive
    files when resuming an extraction run.
    """

    resolved_folder: str | None = folder_id.strip() if folder_id else None
    resolved_files: list[str] = [
        file_id.strip()
        for file_id in file_ids or []
        if file_id and file_id.strip()
    ]
    configured_skip_file_ids = _normalise_identifier_set(skip_file_ids)

    if share_link:
        target_type, identifier = parse_drive_share_link(share_link)
        if target_type == "folder":
            if resolved_folder and resolved_folder != identifier:
                raise ValueError("Conflicting folder identifiers provided")
            resolved_folder = identifier
        else:
            if identifier not in resolved_files:
                resolved_files.append(identifier)

    if not resolved_folder and not resolved_files:
        raise ValueError("At least one folder_id, share_link, or file_id must be provided")

    if client_factory is None:
        def factory() -> GoogleDriveClient:
            return GoogleDriveClient(api_key=api_key, access_token=access_token)
    else:
        factory = client_factory

    base_extractor = pdf_text_extractor or (
        lambda payload, metadata: _default_pdf_text_extractor(
            payload,
            file_name=str(metadata.get("name", "unknown")),
        )
    )

    docx_extractor = None
    if include_docx:
        docx_extractor = docx_text_extractor or (
            lambda payload, metadata: _default_docx_text_extractor(
                payload,
                file_name=str(metadata.get("name", "unknown")),
            )
        )

    languages = _normalise_ocr_languages(ocr_languages) if enable_ocr else None
    dpi = max(int(ocr_dpi or 0), 72)

    def _full_text_extractor(payload: bytes, metadata: Mapping[str, object]) -> str:
        text_result: object | None = None
        try:
            text_result = base_extractor(payload, metadata)
        except Exception as base_error:
            if not enable_ocr:
                raise
            try:
                extracted = _ocr_pdf_text_extractor(
                    payload,
                    file_name=str(metadata.get("name", "unknown")),
                    languages=languages,
                    dpi=dpi,
                )
                return extracted
            except Exception as ocr_error:  # pragma: no cover - fallback path
                raise RuntimeError(
                    "Both direct text extraction and OCR failed for PDF "
                    f"'{metadata.get('name', 'unknown')}'. Direct extraction error: {base_error}"
                ) from ocr_error
        if isinstance(text_result, (list, tuple)):
            combined = "\n\n".join(str(part).strip() for part in text_result if str(part).strip())
            if combined:
                return combined
            text_result = ""
        if isinstance(text_result, bytes):
            text_result = text_result.decode("utf-8", errors="ignore")
        text = str(text_result)
        if text.strip():
            return text
        if not enable_ocr:
            return text
        try:
            extracted = _ocr_pdf_text_extractor(
                payload,
                file_name=str(metadata.get("name", "unknown")),
                languages=languages,
                dpi=dpi,
            )
            return extracted
        except Exception as ocr_error:  # pragma: no cover - fallback path
            if isinstance(text_result, str) and text_result:
                return text_result
            raise RuntimeError(
                "Both direct text extraction and OCR failed for PDF "
                f"'{metadata.get('name', 'unknown')}'"
            ) from ocr_error

    page_chunk_size = _coerce_positive_int(page_batch_size)

    default_tags = tuple(tags or ())
    if enable_ocr and "ocr" not in default_tags:
        default_tags = default_tags + ("ocr",)
    configured_batch_size = _coerce_positive_int(batch_size) or 10

    if include_docx:
        allowed_mime_types: tuple[str, ...] = (_PDF_MIME_TYPE,) + _DOCX_MIME_TYPES
    else:
        allowed_mime_types = (_PDF_MIME_TYPE,)
    allowed_mime_type_set = set(allowed_mime_types)
    drive_mime_types = allowed_mime_types + (_FOLDER_MIME_TYPE,)

    def _resolve_tags(mime_type: str) -> tuple[str, ...]:
        if mime_type == _PDF_MIME_TYPE:
            if "pdf" in default_tags or not default_tags:
                return default_tags
            return default_tags + ("pdf",)
        if mime_type in _DOCX_MIME_TYPES:
            doc_tags = tuple("docx" if tag == "pdf" else tag for tag in default_tags)
            if "docx" not in doc_tags:
                doc_tags = doc_tags + ("docx",)
            return doc_tags
        return default_tags

    def loader(context: CorpusExtractionContext) -> Iterable[Mapping[str, object]]:
        remaining = context.limit
        client = factory()
        seen_ids: set[str] = set()
        try:
            metadata_batch_size: object | None = context.metadata.get("batch_size")
        except AttributeError:  # pragma: no cover - custom mappings without get
            metadata_batch_size = None
        effective_batch_size = (
            _coerce_positive_int(metadata_batch_size)
            or configured_batch_size
        )
        if effective_batch_size > 1000:
            effective_batch_size = 1000
        skip_file_id_set = set(configured_skip_file_ids)
        try:
            metadata_skip_ids = context.metadata.get("skip_file_ids")
        except AttributeError:  # pragma: no cover - custom mappings without get
            metadata_skip_ids = None
        skip_file_id_set.update(_normalise_identifier_set(metadata_skip_ids))

        def _batched(iterator: Iterable[MutableMapping[str, object]]) -> Iterator[list[MutableMapping[str, object]]]:
            batch: list[MutableMapping[str, object]] = []
            for item in iterator:
                batch.append(item)
                if len(batch) >= effective_batch_size:
                    yield batch
                    batch = []
            if batch:
                yield batch

        def _should_stop() -> bool:
            return remaining is not None and remaining <= 0

        visited_folders: set[str] = set()

        def _process_folder(folder_id: str) -> Iterator[Mapping[str, object]]:
            if not folder_id or folder_id in visited_folders:
                return
            visited_folders.add(folder_id)
            iterator = client.iter_files(
                folder_id=folder_id,
                mime_types=drive_mime_types,
                page_size=effective_batch_size,
            )
            for batch in _batched(iterator):
                if _should_stop():
                    break
                for entry in batch:
                    if _should_stop():
                        break
                    yield from _yield_document(entry, parent_folder_id=folder_id)

        def _yield_document(
            metadata: MutableMapping[str, object],
            *,
            parent_folder_id: str | None,
        ) -> Iterator[Mapping[str, object]]:
            nonlocal remaining
            if _should_stop():
                return
            file_id = str(metadata.get("id") or "").strip()
            if not file_id or file_id in skip_file_id_set or file_id in seen_ids:
                return
            seen_ids.add(file_id)
            mime_type = str(metadata.get("mimeType") or "")
            if mime_type == _FOLDER_MIME_TYPE:
                if file_id:
                    yield from _process_folder(file_id)
                return
            if mime_type and mime_type not in allowed_mime_type_set:
                return
            size_value = _coerce_size(metadata.get("size"))
            if max_file_size is not None and size_value is not None and size_value > max_file_size:
                return
            payload = client.download_file(file_id)
            if max_file_size is not None and len(payload) > max_file_size:
                return
            if mime_type:
                effective_mime = mime_type
            else:
                name = str(metadata.get("name") or "")
                if include_docx and name.lower().endswith(".docx"):
                    effective_mime = _DOCX_MIME_TYPES[0]
                else:
                    effective_mime = _PDF_MIME_TYPE
            document_metadata: dict[str, object] = {
                "file_id": file_id,
                "file_name": metadata.get("name"),
                "mime_type": effective_mime,
            }
            if parent_folder_id:
                document_metadata["parent_folder_id"] = parent_folder_id
            if resolved_folder:
                document_metadata["source_folder_id"] = resolved_folder
            if resolved_files:
                document_metadata["source_file_ids"] = tuple(resolved_files)
            if "modifiedTime" in metadata:
                document_metadata["modified_time"] = metadata["modifiedTime"]
            if "md5Checksum" in metadata:
                document_metadata["md5_checksum"] = metadata["md5Checksum"]
            if "webViewLink" in metadata:
                document_metadata["web_view_link"] = metadata["webViewLink"]
            if size_value is not None:
                document_metadata["size"] = size_value
            parents_field = metadata.get("parents")
            if isinstance(parents_field, str):
                parent_value = parents_field.strip()
                if parent_value:
                    document_metadata["parents"] = (parent_value,)
            elif isinstance(parents_field, Sequence):
                parent_ids = tuple(
                    str(candidate).strip()
                    for candidate in parents_field
                    if str(candidate).strip()
                )
                if parent_ids:
                    document_metadata["parents"] = parent_ids

            tags_for_document = _resolve_tags(effective_mime)

            if page_chunk_size is not None and effective_mime == _PDF_MIME_TYPE:
                if pdf_text_extractor is None:
                    page_texts = _extract_pdf_page_texts(
                        payload,
                        file_name=str(metadata.get("name", "unknown")),
                        enable_ocr=enable_ocr,
                        languages=languages,
                        dpi=dpi,
                    )
                else:
                    result = base_extractor(payload, metadata)
                    if isinstance(result, (list, tuple)):
                        page_texts = [str(part).strip() for part in result]
                    else:
                        raise RuntimeError(
                            "Custom pdf_text_extractor must return a sequence of page texts "
                            "when page_batch_size is provided"
                        )

                total_pages = len(page_texts)
                if total_pages == 0:
                    raise RuntimeError(
                        f"No pages discovered while extracting PDF '{metadata.get('name', 'unknown')}'"
                    )
                total_batches = (total_pages + page_chunk_size - 1) // page_chunk_size
                produced = False
                for batch_number, start_index in enumerate(
                    range(0, total_pages, page_chunk_size), start=1
                ):
                    if _should_stop():
                        break
                    end_index = min(start_index + page_chunk_size, total_pages)
                    chunk_parts = [
                        part.strip() for part in page_texts[start_index:end_index] if part.strip()
                    ]
                    if not chunk_parts:
                        continue
                    content = "\n\n".join(chunk_parts)
                    chunk_metadata = dict(document_metadata)
                    chunk_metadata.update(
                        page_start=start_index + 1,
                        page_end=end_index,
                        page_count=end_index - start_index,
                        total_pages=total_pages,
                        batch_number=batch_number,
                        total_batches=total_batches,
                        pages_per_batch=page_chunk_size,
                    )
                    yield {
                        "identifier": (
                            f"google-drive-{file_id}-p{start_index + 1:04d}-p{end_index:04d}"
                        ),
                        "content": content,
                        "metadata": chunk_metadata,
                        "tags": tags_for_document,
                    }
                    produced = True
                    if remaining is not None:
                        remaining -= 1
                        if remaining <= 0:
                            break
                if not produced:
                    raise RuntimeError(
                        f"No extractable text found in PDF '{metadata.get('name', 'unknown')}'"
                    )
            else:
                if effective_mime in _DOCX_MIME_TYPES:
                    if docx_extractor is None:
                        return
                    text = docx_extractor(payload, metadata)
                else:
                    text = _full_text_extractor(payload, metadata)
                yield {
                    "identifier": f"google-drive-{file_id}",
                    "content": text,
                    "metadata": document_metadata,
                    "tags": tags_for_document,
                }
                if remaining is not None:
                    remaining -= 1

        if resolved_folder:
            yield from _process_folder(resolved_folder)

        if not _should_stop() and resolved_files:
            metadata_batch: list[MutableMapping[str, object]] = []
            for file_id in resolved_files:
                if _should_stop():
                    break
                metadata = client.get_file_metadata(
                    file_id,
                    fields="id, name, mimeType, modifiedTime, size, md5Checksum, webViewLink",
                )
                if not isinstance(metadata, MutableMapping):
                    continue
                metadata_batch.append(metadata)
                if len(metadata_batch) >= effective_batch_size:
                    for entry in metadata_batch:
                        if _should_stop():
                            break
                        yield from _yield_document(entry, parent_folder_id=None)
                    metadata_batch.clear()
            if metadata_batch and not _should_stop():
                for entry in metadata_batch:
                    if _should_stop():
                        break
                    yield from _yield_document(entry, parent_folder_id=None)

    return loader

